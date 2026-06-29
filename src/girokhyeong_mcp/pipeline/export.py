# -*- coding: utf-8 -*-
"""제출본 변환 — 마크다운 답안 → docx (가인/독립판 서식). 최종 HWP/HWPX 는 한글에서 변환.

python-docx 가 있으면 동작. 면수 실측은 한컴/Word 의존이라 MCP 범위 밖(메시지로 안내).
"""
from __future__ import annotations

import re
from pathlib import Path

from .. import config


def available() -> bool:
    try:
        import docx  # noqa
        return True
    except Exception:
        return False


def md_to_docx(in_md: str, out_docx: str, *, body_pt: int | None = None,
               line_pct: int | None = None, title: str | None = None) -> dict:
    """마크다운 답안 → docx. 서식 기본값은 config.FORMAT_DEFAULTS(바탕 11pt·160%·양쪽맞춤·A4·여백 20mm)."""
    try:
        import docx
        from docx.shared import Pt, Mm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    except Exception as e:
        return {"error": f"python-docx 미설치: {e} (pip install python-docx)"}

    fd = config.FORMAT_DEFAULTS
    body_pt = body_pt or fd["body_pt"]
    line_pct = line_pct or fd["line_pct"]

    src = Path(in_md)
    text = src.read_text(encoding="utf-8") if src.exists() else in_md
    # 프론트매터 제거
    text = re.sub(r"^---\n.*?\n---\n", "", text, count=1, flags=re.S)

    doc = docx.Document()
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(section, attr, Mm(fd["margin_mm"]))

    normal = doc.styles["Normal"]
    normal.font.name = fd["font"]
    normal.font.size = Pt(body_pt)
    try:
        normal.element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", fd["font"])
    except Exception:
        pass

    def add(text_line, *, pt=body_pt, bold=False, align=None, center=False):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = line_pct / 100.0
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "justify" or fd["align"] == "justify":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text_line)
        run.bold = bold
        run.font.size = Pt(pt)
        run.font.name = fd["font"]
        return p

    if title:
        add(title, pt=fd["title_pt"], bold=True, center=True)
        doc.add_paragraph()

    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            add(m.group(2), pt=fd["title_pt"] if level == 1 else body_pt + 1,
                bold=True, center=(level == 1))
        else:
            add(line)

    Path(out_docx).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)
    return {"saved": out_docx, "body_pt": body_pt, "line_pct": line_pct,
            "note": "면수 실측·HWP/HWPX 변환은 한글에서. docx↔한글 페이지수 차이 가능 — 캡−1쪽 여유 권장."}
